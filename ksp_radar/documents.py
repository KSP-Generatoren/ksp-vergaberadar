"""Vergabeunterlagen holen und lesbar machen.

Rechtsgrundlage: Paragraf 41 Abs. 1 VgV (oberschwellig) und Paragraf 29 UVgO
(unterschwellig) verlangen wortgleich, dass der Auftraggeber in der
Bekanntmachung eine elektronische Adresse angibt, unter der die
Vergabeunterlagen "unentgeltlich, uneingeschraenkt, vollstaendig und direkt"
abrufbar sind. "Direkt" ist in Rechtsprechung und Literatur als "ohne
Registrierung" ausgelegt.

Diese Adresse steht als strukturiertes Feld in der Bekanntmachung:
eForms BT-15 "Documents URL". Deshalb ist der Weg von der Bekanntmachung
zum Volltext der Unterlagen durchgaengig automatisierbar.
"""
from __future__ import annotations

import hashlib
import io
import logging
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path

import httpx

from .config import USER_AGENT
from .models import DocumentSnapshot

log = logging.getLogger(__name__)

TEXT_SUFFIXES = {".pdf", ".docx", ".doc", ".txt", ".rtf", ".xml", ".html"}
MAX_BYTES = 120 * 1024 * 1024


def fetch_documents(notice_uid: str, url: str, cache_dir: str = "./data/docs") -> DocumentSnapshot | None:
    """Laedt die Vergabeunterlagen und erstellt einen Abzug.

    Der Abzug ist die Grundlage des Aenderungsmonitors: gleiche URL,
    spaeterer Zeitpunkt, anderer Hash -> die Vergabestelle hat nachgebessert.
    """
    if not url:
        return None

    out = Path(cache_dir) / notice_uid
    out.mkdir(parents=True, exist_ok=True)

    try:
        with httpx.Client(timeout=180, headers={"User-Agent": USER_AGENT},
                          follow_redirects=True) as client:
            r = client.get(url)
            r.raise_for_status()
            payload = r.content[:MAX_BYTES]
    except httpx.HTTPError as exc:
        log.warning("Unterlagen %s nicht abrufbar: %s", url, exc)
        return None

    files: dict[str, str] = {}
    texts: list[str] = []

    if _looks_like_zip(payload):
        with zipfile.ZipFile(io.BytesIO(payload)) as zf:
            for info in zf.infolist():
                if info.is_dir():
                    continue
                data = zf.read(info)
                files[info.filename] = hashlib.sha256(data).hexdigest()[:16]
                (out / Path(info.filename).name).write_bytes(data)
                if Path(info.filename).suffix.lower() in TEXT_SUFFIXES:
                    texts.append(f"\n\n=== {info.filename} ===\n" + extract_text(info.filename, data))
    else:
        name = _filename_from_url(url)
        files[name] = hashlib.sha256(payload).hexdigest()[:16]
        (out / name).write_bytes(payload)
        texts.append(extract_text(name, payload))

    return DocumentSnapshot(
        notice_uid=notice_uid,
        fetched_at=datetime.now(timezone.utc),
        files=files,
        text="\n".join(texts),
    )


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    try:
        if suffix == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            return "\n".join((p.extract_text() or "") for p in reader.pages)
        if suffix == ".docx":
            import docx
            d = docx.Document(io.BytesIO(data))
            parts = [p.text for p in d.paragraphs]
            for table in d.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return "\n".join(parts)
        if suffix in {".txt", ".xml", ".html", ".rtf"}:
            return data.decode("utf-8", errors="replace")
    except Exception as exc:                          # noqa: BLE001
        log.warning("Textextraktion %s fehlgeschlagen: %s", filename, exc)
    return ""


def _looks_like_zip(data: bytes) -> bool:
    return data[:2] == b"PK"


def _filename_from_url(url: str) -> str:
    name = re.sub(r"[?#].*$", "", url).rstrip("/").rsplit("/", 1)[-1]
    return name or "download.bin"
